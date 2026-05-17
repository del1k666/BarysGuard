"""Convert diploma.md to diploma.docx with academic formatting."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH   = Path(__file__).parent / "diploma.md"
DOCX_PATH = Path(__file__).parent / "diploma.docx"

FONT_MAIN = "Times New Roman"
FONT_CODE = "Courier New"
SIZE_MAIN = 14
SIZE_H1   = 16
SIZE_H2   = 14
SIZE_H3   = 13

# ── helpers ─────────────────────────────────────────────────────────────────

def set_paragraph_format(para, left_cm=0, first_cm=1.25, space_before=0,
                          space_after=6, line_spacing=1.5, align=None):
    pf = para.paragraph_format
    pf.left_indent       = Cm(left_cm)
    pf.first_line_indent = Cm(first_cm)
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.line_spacing      = line_spacing  # float multiplier (1.5 = полуторный)
    if align:
        pf.alignment = align


def set_run_font(run, bold=False, italic=False, size=SIZE_MAIN, font=FONT_MAIN, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    font)
    rFonts.set(qn("w:hAnsi"),    font)
    rFonts.set(qn("w:cs"),       font)
    rFonts.set(qn("w:eastAsia"), font)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_run_inline(para, text, bold=False, italic=False, code=False):
    """Add a run respecting **bold**, *italic*, `code` inline markup."""
    run = para.add_run(text)
    if code:
        set_run_font(run, bold=False, italic=False, size=11, font=FONT_CODE)
    else:
        set_run_font(run, bold=bold, italic=italic)
    return run


def parse_inline(para, text):
    """Parse inline markdown (**bold**, *italic*, `code`) into runs."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            add_run_inline(para, text[pos:m.start()])
        full = m.group(0)
        if full.startswith("**"):
            add_run_inline(para, m.group(2), bold=True)
        elif full.startswith("*"):
            add_run_inline(para, m.group(3), italic=True)
        else:
            add_run_inline(para, m.group(4), code=True)
        pos = m.end()
    if pos < len(text):
        add_run_inline(para, text[pos:])


def shade_cell(cell, fill="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)


# ── main conversion ──────────────────────────────────────────────────────────

def convert():
    doc = Document()

    # Page margins (GOST-ish): left 3cm, rest 2cm
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # ── code block ──────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                para = doc.add_paragraph()
                para.paragraph_format.left_indent       = Cm(1)
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.space_before      = Pt(4)
                para.paragraph_format.space_after       = Pt(4)
                para.paragraph_format.line_spacing      = 1.0
                run = para.add_run("\n".join(code_lines))
                set_run_font(run, size=10, font=FONT_CODE)
                # grey background via shading on the paragraph's run (approximation)
                pPr = para._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "F4F4F4")
                pPr.append(shd)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── table (| … | rows) ──────────────────────────────────────────────
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # filter separator rows
            rows = [r for r in table_lines if not re.match(r"^\s*\|[-|: ]+\|\s*$", r)]
            if not rows:
                continue

            parsed_rows = []
            for r in rows:
                cells = [c.strip() for c in r.strip().strip("|").split("|")]
                parsed_rows.append(cells)

            if not parsed_rows:
                continue
            ncols = max(len(r) for r in parsed_rows)
            # pad rows
            parsed_rows = [r + [""] * (ncols - len(r)) for r in parsed_rows]

            tbl = doc.add_table(rows=len(parsed_rows), cols=ncols)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            for ri, row_data in enumerate(parsed_rows):
                for ci, cell_text in enumerate(row_data):
                    cell = tbl.cell(ri, ci)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    parse_inline(p, cell_text)
                    for run in p.runs:
                        run.font.name = FONT_MAIN
                        run.font.size = Pt(11)
                        if ri == 0:
                            run.bold = True
                    if ri == 0:
                        shade_cell(cell, "BFBFBF")
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
            doc.add_paragraph()  # spacing after table
            continue

        stripped = line.strip()

        # ── blank line ───────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── headings ─────────────────────────────────────────────────────────
        h_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if h_match:
            level  = len(h_match.group(1))
            text   = h_match.group(2)
            style  = f"Heading {level}"
            para   = doc.add_paragraph(style=style)
            para.clear()
            run = para.add_run(text)
            if level == 1:
                set_run_font(run, bold=True,  size=SIZE_H1)
                para.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(24)
                para.paragraph_format.space_after  = Pt(12)
            elif level == 2:
                set_run_font(run, bold=True,  size=SIZE_H2)
                para.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(18)
                para.paragraph_format.space_after  = Pt(6)
            else:
                set_run_font(run, bold=True,  size=SIZE_H3)
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after  = Pt(4)
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.line_spacing      = 1.2
            i += 1
            continue

        # ── horizontal rule ──────────────────────────────────────────────────
        if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
            doc.add_paragraph()
            i += 1
            continue

        # ── bullet list ──────────────────────────────────────────────────────
        bl_match = re.match(r"^[-*+]\s+(.*)", stripped)
        if bl_match:
            para = doc.add_paragraph(style="List Bullet")
            para.clear()
            parse_inline(para, bl_match.group(1))
            for run in para.runs:
                run.font.name = FONT_MAIN
                run.font.size = Pt(SIZE_MAIN)
            para.paragraph_format.left_indent       = Cm(1)
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before      = Pt(0)
            para.paragraph_format.space_after       = Pt(3)
            para.paragraph_format.line_spacing      = 1.15
            i += 1
            continue

        # ── numbered list ────────────────────────────────────────────────────
        nl_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if nl_match:
            para = doc.add_paragraph(style="List Number")
            para.clear()
            parse_inline(para, nl_match.group(1))
            for run in para.runs:
                run.font.name = FONT_MAIN
                run.font.size = Pt(SIZE_MAIN)
            para.paragraph_format.left_indent       = Cm(1)
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before      = Pt(0)
            para.paragraph_format.space_after       = Pt(3)
            para.paragraph_format.line_spacing      = 1.15
            i += 1
            continue

        # ── blockquote / figure captions ─────────────────────────────────────
        bq_match = re.match(r"^>\s*(.*)", stripped)
        if bq_match:
            para = doc.add_paragraph()
            parse_inline(para, bq_match.group(1))
            for run in para.runs:
                set_run_font(run, italic=True, size=SIZE_MAIN - 1)
            para.paragraph_format.left_indent       = Cm(1.5)
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before      = Pt(4)
            para.paragraph_format.space_after       = Pt(4)
            i += 1
            continue

        # ── regular paragraph ────────────────────────────────────────────────
        para = doc.add_paragraph()
        parse_inline(para, stripped)
        for run in para.runs:
            if run.font.name not in (FONT_CODE,):
                run.font.name = FONT_MAIN
                if run.font.size is None:
                    run.font.size = Pt(SIZE_MAIN)
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        i += 1

    doc.save(str(DOCX_PATH))
    print(f"Saved: {DOCX_PATH}")
    print(f"Size:  {DOCX_PATH.stat().st_size // 1024} KB")


if __name__ == "__main__":
    convert()
